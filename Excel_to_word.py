import docx
from docx import Document
import pandas as pd


document = Document()
# 数据导入
video_list = pd.read_excel('./data/video_list.xlsx')
speech_text = pd.read_excel('./data/speech_text.xlsx')

# video_list.info()
# speech_text.info()
# 通过观察发现文件的AwemeId和VideoId为int类型，但是更希望是字符串以便后续的操作，所以进行转换
video_list['AwemeId'] = video_list['AwemeId'].astype(str)
speech_text['VideoId'] = speech_text['VideoId'].astype(str)

# 两表分开但是希望整合在同一DataFrame下
merge = pd.merge(video_list, speech_text, how='inner' , left_on='AwemeId', right_on='VideoId')
# print(merge.head(10))

#将数据存入doc文件当中
for i in range(len(merge)):
    if merge.iloc[i]['品牌'] != merge.iloc[i - 1]['品牌'] or i == 0:
        document.add_heading(merge['品牌'][i], level=1)
    document.add_heading(merge['视频标题'][i], level=2)
    document.add_paragraph(f'达人昵称{merge.iloc[i]["BloggerName"]}')

    # 创建包含超链接的段落
    para = document.add_paragraph('视频地址：')
    url = f'https://www.douyin.com/video/{merge.iloc[i]["AwemeId"]}'
    # 添加超链接文本
    run = para.add_run(url)
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)  # 蓝色
    run.font.underline = True  # 添加下划线

    document.add_paragraph(merge['视频文案'][i])

document.save('./output/demo01.docx')
